"""
Internal document indexing module: Load, chunk, embed, and index internal documents
"""
import logging
from pathlib import Path
from typing import List, Dict
import requests

import chromadb
from chromadb.config import Settings

from app.config import (
    INTERNAL_DATA_DIR, CHROMA_DIR, EMBED_MODEL, OLLAMA_BASE_URL,
    CHUNK_SIZE, CHUNK_OVERLAP
)
from app.utils import chunk_text, logger

logger = logging.getLogger(__name__)


def get_embedding(text: str) -> List[float]:
    """
    Get embedding for text using Ollama
    
    Args:
        text: Text to embed
    
    Returns:
        Embedding vector
    """
    try:
        url = f"{OLLAMA_BASE_URL}/api/embeddings"
        payload = {
            "model": EMBED_MODEL,
            "prompt": text
        }
        
        response = requests.post(url, json=payload, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        embedding = result.get("embedding", [])
        
        if not embedding:
            raise ValueError("Ollama returned empty embedding")
        
        return embedding
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to get embedding: {e}")
        raise ConnectionError(f"Failed to connect to Ollama for embeddings: {e}")


def load_documents(data_dir: Path) -> List[Dict[str, str]]:
    """
    Load documents from data directory (PDF, TXT, MD, DOCX)
    
    Args:
        data_dir: Directory containing internal documents
    
    Returns:
        List of documents with metadata
    """
    documents = []
    
    # Supported file extensions
    extensions = {'.pdf', '.txt', '.md', '.docx'}
    
    if not data_dir.exists():
        logger.warning(f"Data directory does not exist: {data_dir}")
        return documents
    
    # First try direct files in the directory
    for file_path in data_dir.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            try:
                if file_path.suffix.lower() == '.pdf':
                    from app.pdf_extract import extract_full_text
                    text = extract_full_text(file_path, max_pages=1000)  # Load all pages for internal docs
                elif file_path.suffix.lower() in ['.txt', '.md']:
                    text = file_path.read_text(encoding='utf-8')
                elif file_path.suffix.lower() == '.docx':
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        text = '\n'.join([para.text for para in doc.paragraphs])
                    except ImportError:
                        logger.warning(f"python-docx not installed, skipping {file_path}")
                        continue
                    except Exception as e:
                        logger.warning(f"Failed to read DOCX {file_path}: {e}")
                        continue
                else:
                    continue
                
                if text.strip():
                    documents.append({
                        'doc_id': file_path.stem,
                        'doc_title': file_path.name,
                        'doc_path': str(file_path),
                        'text': text
                    })
                    logger.info(f"Loaded document: {file_path.name} ({len(text)} chars)")
                    
            except ImportError as import_err:
                error_msg = str(import_err)
                if "pypdf" in error_msg.lower() or "pdfplumber" in error_msg.lower():
                    logger.error(f"PDF library not installed. Please install: pip install pypdf pdfplumber")
                    logger.error(f"Or install all dependencies: pip install -r requirements.txt")
                else:
                    logger.error(f"Missing required library for {file_path}: {import_err}")
                    logger.error("Please install required dependencies: pip install -r requirements.txt")
                raise
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                # Don't raise here, continue with other files
                continue
    
    # Also try rglob for subdirectories
    for file_path in data_dir.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            # Skip if already processed
            if any(doc['doc_path'] == str(file_path) for doc in documents):
                continue
            try:
                if file_path.suffix.lower() == '.pdf':
                    from app.pdf_extract import extract_full_text
                    text = extract_full_text(file_path, max_pages=1000)  # Load all pages for internal docs
                elif file_path.suffix.lower() in ['.txt', '.md']:
                    text = file_path.read_text(encoding='utf-8')
                elif file_path.suffix.lower() == '.docx':
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        text = '\n'.join([para.text for para in doc.paragraphs])
                    except ImportError:
                        logger.warning(f"python-docx not installed, skipping {file_path}")
                        continue
                    except Exception as e:
                        logger.warning(f"Failed to read DOCX {file_path}: {e}")
                        continue
                else:
                    continue
                
                if text.strip():
                    documents.append({
                        'doc_id': file_path.stem,
                        'doc_title': file_path.name,
                        'doc_path': str(file_path),
                        'text': text
                    })
                    logger.info(f"Loaded document: {file_path.name} ({len(text)} chars)")
                    
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
    
    return documents


def index_internal_documents():
    """
    Main function to index all internal documents
    Processes company/EDU/company_data.pdf and stores in vector DB
    """
    logger.info("Starting internal document indexing")
    logger.info(f"Looking for documents in: {INTERNAL_DATA_DIR}")
    # INTERNAL_DATA_DIR is already resolved in config.py, but ensure it's absolute
    internal_dir = Path(INTERNAL_DATA_DIR).resolve()
    logger.info(f"Absolute path: {internal_dir}")
    logger.info(f"Path exists: {internal_dir.exists()}")
    
    # Ensure we use resolved absolute path
    internal_dir = Path(INTERNAL_DATA_DIR).resolve()
    
    # Verify directory exists
    if not internal_dir.exists():
        logger.error(f"Directory does not exist: {internal_dir}")
        # Try to find the correct path
        possible_paths = [
            Path(__file__).parent.parent.parent / "company" / "EDU",  # From rag_demo root
            Path(__file__).parent.parent.parent.parent / "company" / "EDU",  # Alternative
        ]
        for possible in possible_paths:
            resolved_possible = possible.resolve()
            logger.info(f"Checking possible path: {resolved_possible} (exists: {resolved_possible.exists()})")
            if resolved_possible.exists():
                logger.warning(f"Found directory at: {resolved_possible}")
                internal_dir = resolved_possible
                break
        else:
            raise FileNotFoundError(f"Internal data directory not found: {internal_dir}")
    
    # Load documents
    documents = load_documents(internal_dir)
    
    if not documents:
        logger.warning(f"No documents found in {internal_dir}")
        logger.warning(f"Absolute path: {internal_dir}")
        logger.warning(f"Expected file: {internal_dir / 'company_data.pdf'}")
        # List actual files in directory
        if internal_dir.exists():
            actual_files = list(internal_dir.glob('*'))
            logger.warning(f"Actual files in directory: {[f.name for f in actual_files]}")
            # Check if PDF extraction failed
            pdf_files = list(internal_dir.glob('*.pdf'))
            if pdf_files:
                logger.error(f"PDF files found but extraction failed. Check if PDF libraries are installed:")
                logger.error(f"  pip install pypdf pdfplumber")
        raise ValueError(f"No documents found in {internal_dir}. Please check if PDF libraries are installed: pip install pypdf pdfplumber")
    
    logger.info(f"Found {len(documents)} document(s) to index:")
    for doc in documents:
        logger.info(f"  - {doc['doc_title']} ({len(doc['text'])} characters)")
    
    # Initialize ChromaDB
    client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=Settings(anonymized_telemetry=False)
    )
    
    # Get or create collection
    collection = client.get_or_create_collection(
        name="internal_documents",
        metadata={"description": "Internal company documents for rebuttal"}
    )
    
    # Check if collection already has data - if yes, skip indexing
    if collection.count() > 0:
        logger.info(f"Collection already exists with {collection.count()} items. Skipping indexing.")
        return
    
    # Process each document
    all_chunks = []
    all_metadatas = []
    all_ids = []
    
    for doc in documents:
        # Chunk the document
        chunks = chunk_text(doc['text'], chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        
        logger.info(f"Chunked {doc['doc_title']} into {len(chunks)} chunks")
        
        # Create embeddings and metadata for each chunk
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc['doc_id']}_chunk_{i}"
            
            all_chunks.append(chunk)
            all_metadatas.append({
                'doc_id': doc['doc_id'],
                'doc_title': doc['doc_title'],
                'doc_path': doc['doc_path'],
                'chunk_id': chunk_id,
                'chunk_index': i
            })
            all_ids.append(chunk_id)
    
    # Batch embed and add to collection
    logger.info(f"Generating embeddings for {len(all_chunks)} chunks...")
    batch_size = 10
    
    for i in range(0, len(all_chunks), batch_size):
        batch_chunks = all_chunks[i:i+batch_size]
        batch_ids = all_ids[i:i+batch_size]
        batch_metadatas = all_metadatas[i:i+batch_size]
        
        # Get embeddings
        embeddings = []
        for chunk in batch_chunks:
            try:
                embedding = get_embedding(chunk)
                embeddings.append(embedding)
            except Exception as e:
                logger.error(f"Failed to embed chunk {batch_ids[len(embeddings)]}: {e}")
                # Use zero vector as fallback (not ideal, but allows processing to continue)
                embeddings.append([0.0] * 768)
        
        # Add to collection
        collection.add(
            ids=batch_ids,
            embeddings=embeddings,
            documents=batch_chunks,
            metadatas=batch_metadatas
        )
        
        logger.info(f"Indexed batch {i//batch_size + 1}/{(len(all_chunks) + batch_size - 1)//batch_size}")
    
    logger.info(f"Successfully indexed {collection.count()} chunks from {len(documents)} documents")
    logger.info(f"ChromaDB collection saved to {CHROMA_DIR}")


if __name__ == "__main__":
    index_internal_documents()
